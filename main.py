import docx
from docx import Document
from docx.shared import RGBColor
import re
import xlwings
from tkinter import *
from tkinter import filedialog
import PIL.Image
import PIL.ImageTk
import os
import subprocess, platform

# DER and TBV are not valid tags
docRelation = {"HRD": ("HRS"), "HRS": ("PRS"), "PRS": ("URS", "RISK"), "HTR": ("HTP"), "HTP": ("HRD", "HRS"), \
               "SDS": ("BOLUS", "ACE", "AID"), "ACE": ("PRS"), "BOLUS": ("PRS"), "AID": ("PRS"), \
               "SVAL": ("BOLUS", "ACE", "AID"), "SVATR": ("SVAL"), "UT": ("UNIT"),
               "INS": ("UNIT")}  # to be created by the GUI

docFile = {"HRD": "HDS_new_pump.docx", "HRS": "HRS_new_pump.docx", "HTP": "HTP_new_pump.docx",
           "HTR": "HTR_new_pump.docx", \
           "PRS": "PRS_new_pump.docx", "RISK": "RiskAnalysis_Pump.docx", "SDS": "SDS_New_pump_x04.docx", \
           "ACE": "SRS_ACE_Pump_X01.docx", "BOLUS": "SRS_BolusCalc_Pump_X04.docx",
           "SRS": "SRS_DosingAlgorithm_X03.docx", \
           "SVAL": "SVaP_new_pump.docx", "SVATR": "SVaTR_new_pump.docx", "UT": "SVeTR_new_pump.docx",
           "URS": "URS_new_pump.docx"}

# docFileEmpty = {}  # declares new empty dictionary for Gui and doesn't interfere with docFile dictionary

docFileList = list(docFile.keys())  # This is a list of all main tags found in each document
# print(docFileList)
# docFileEmptyList = list(docFileEmpty.keys())
parentTagList = list(docRelation.values())

report1 = Document()  # create word document
paragraph = report1.add_paragraph()
report1.save('report1.docx')

uniqueValidTagList = []  # This is the valid child tag list
for tag in parentTagList:
    if type(tag) is tuple:  # if a tuple is found, convert to a list and add to the list
        uniqueValidTagList.extend(list(tag))
    else:
        uniqueValidTagList.append(tag)  # if not a tuple simply append to the list
uniqueTagList = (list(set(uniqueValidTagList)))  # set() strips out all redundant tags

# filePath = "C:/Users/Willi/Desktop/Tom's example docs/"
filePath = "/Users/adrian/Desktop/SampledocsTandem/"


def GetText(filename):  # Opens the document and places each paragraph into a list
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    fullText = [ele for ele in fullText if ele.strip()]  # Eliminates empty paragraphs
    return fullText


def GetParentTags():  # Returns only valid parent tags
    for tag in docFileList:  # Tags are used to open the corresponding file
        textList = GetText(filePath + docFile[tag])
        index = 0
        ind = []
        for t in textList:
            if tag == "BOLUS" or tag == "ACE":
                if re.search('.*[:\s]' + "SRS" + '[:\s]', t):
                    ind.append(index)
                    tt = t
                    y = re.findall('\S*[:\s]' + "SRS" + '[:\s]\S*', t)
                    # red = paragraph.add_run(y)
                    # paragraph.add_run("\n\n")
                    # red.bold = True
                    # red.font.color.rgb = RGBColor(255, 0, 0)

                    # print(y[0])
                    qe.append(y[0])  # adds to parent tag list
                index = index + 1
            # print(ind)
            else:
                if re.search('.*[:\s]' + re.escape(tag) + '[:\s]', t):
                    ind.append(index)
                    tt = t
                    y = re.findall('\S*[:\s]' + re.escape(tag) + '[:\s]\S*', t)
                    # red = paragraph.add_run(y)
                    # paragraph.add_run("\n\n")
                    # red.bold = True
                    # red.font.color.rgb = RGBColor(255, 0, 0)
                    # print(y[0])
                    qe.append(y[0])  # adds to parent tag list
                index = index + 1

            # print(ind)


def GetChildTags():  # Returns only valid child tags
    for tag in docFileList:  # Tags are used to open the corresponding file
        textList = GetText(filePath + docFile[tag])
        index = 0
        ind = []
        for t in textList:
            if tag == "BOLUS" or tag == "ACE":
                if re.search('.*[:\s]' + "SRS" + '[:\s]', t):
                    ind.append(index)
                    tt = t
                    y = re.findall('\[.+\]', t)
                    if len(y) != 0:
                        # green = paragraph.add_run(y[0])
                        # paragraph.add_run("\n\n")
                        # green.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
                        # green.bold = True
                        # print(y[0])
                        qa.append(y[0])  # adds to child tag list
                index = index + 1
                # print(ind)
            else:
                if re.search('.*[:\s]' + re.escape(tag) + '[:\s]', t):
                    ind.append(index)
                    tt = t
                    y = re.findall('\[.+\]', t)
                    if len(y) != 0:
                        # green = paragraph.add_run(y[0])
                        # paragraph.add_run("\n\n")
                        # green.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
                        # green.bold = True
                        # print(y[0])
                        qa.append(y[0])  # adds to child tag list
                index = index + 1

            # print(ind)


def GetOrphanTags():
    for tag in docFileList:  # Tags are used to open the corresponding file
        textList = GetText(filePath + docFile[tag])
        index = 0
        ind = []
        for t in textList:
            # y = re.findall('[\s\]]\[.+\][\[\s]', t)
            y = re.findall('\[.+\]', t)
            if len(y) != 0:
                # green = paragraph.add_run(y[0])
                # paragraph.add_run("\n\n")
                # green.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
                # green.bold = True

                # print(y[0])
                q.append(y[0])  # adds to orphan tag list
                index = index + 1

                # print(ind)

            # else:
            #    if re.search('.*[:\s]' + re.escape(tag) + '[:\s]', t):
            #        ind.append(index)
            #        tt = t
            #        y = re.findall('\s\[.+\]\s', t)
            #        if len(y) != 0:
            #            green = paragraph.add_run(y[0])
            #            paragraph.add_run("\n\n")
            #            green.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
            #            green.bold = True
            #            # print(y[0])
            #    index = index + 1
            # print(ind)


qe = []  # list of parent tags
qa = []  # List of child tags
q = []  # List for orphans
# runner2 = paragraph.add_run("\n\nParent tag/tags\n\n")
# runner2.bold = True                              #make it bold
GetParentTags()

# runner2 = paragraph.add_run("\n\nChild tag/tags\n\n")
# runner2.bold = True
GetChildTags()
# print(qe)  # prints Parent tags
# print(qa)  # prints Child tags
GetOrphanTags()
# print(q)
# num = 0
# while qe and qa:
# green = paragraph.add_run(qe[0] + "    ")
# paragraph.add_run("\n\n")
# green.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
# green.bold = True
# qe.remove(qe[0])
# red = paragraph.add_run(qa[0])
# paragraph.add_run("\n\n")
# red.bold = True
# red.font.color.rgb = RGBColor(255, 0, 0)
# qa.remove(qa[0])

# while qe:
# green = paragraph.add_run(qe[0])
# paragraph.add_run("\n\n")
# green.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
# green.bold = True
# qe.remove(qe[0])

table = report1.add_table(rows=1, cols=2)

# Adding heading in the 1st row of the table
row = table.rows[0].cells
row[0].text = 'Parent Tag'
row[1].text = 'Child Tag/Tags'

# Adding style to a table
table.style = 'Colorful List'

while qe and qa:
    row = table.add_row().cells  # Adding a row and then adding data in it.
    row[0].text = qe[0]
    #green = paragraph.add_run(sentences[0] + "    ")
    #paragraph.add_run("\n\n")
    #green.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
    #green.bold = True

    #red = paragraph.add_run(child[0])

    # green = paragraph.add_run(sentences[0] + "    ")
    # paragraph.add_run("\n\n")
    # green.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
    # green.bold = True
    # red = paragraph.add_run(child[0])

    noChild = ["PUMP:RISK"]
    if qa:
        if noChild[0] not in qe[0]:
            row[1].text = qa[0]
            qa.remove(qa[0])
            qe.remove(qe[0])
        else:
            row[1].text = " "
        #paragraph.add_run("\n\n")
        #red.bold = True
        #red.font.color.rgb = RGBColor(255, 0, 0)

            # paragraph.add_run("\n\n")
            # red.bold = True
            # red.font.color.rgb = RGBColor(255, 0, 0)
            qe.remove(qe[0])

while qe:
    row = table.add_row().cells  # Adding a row and then adding data in it.
    row[0].text = qe[0]
    # green = paragraph.add_run(sentences[0])
    # paragraph.add_run("\n\n")
    # green.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
    # green.bold = True
    qe.remove(qe[0])

# paragraph.add_run(f)

# runner2 = paragraph.add_run("\n\nOrphanChild tag/tags\n")
# runner2.bold = True
# GetOrphanTags()

report1.save('report1.docx')


# GetOrphanTags()

def readtxt(filename, color: tuple[int, int, int]):
    doc = docx.Document(filename)
    text10 = ""
    fullText = []
    new = []
    for para in doc.paragraphs:

        # Getting the colored words from the doc
        if (getcoloredTxt(para.runs, color)):
            # Concatenating list of runs between the colored text to single a string

            sentence = "".join(r.text for r in para.runs)
            fullText.append(sentence)
            print(sentence)
            text10 = sentence
            # new = (sentence.replace (']', ']\n\n'))
            parent.append("".join(r.text for r in para.runs))

    # print(fullText)
    global filtered_L
    global hasChild
    global fullText2
    filtered_L = [value for value in fullText if "[" not in value]
    hasChild = [value for value in fullText if "[" in value]
    fullText2 = [value for value in fullText]

    # print(filtered_L)
    # print(fullText)
    return fullText, filtered_L, hasChild


def getcoloredTxt(runs, color):
    coloredWords, word = [], ""
    for run in runs:
        if run.font.color.rgb == RGBColor(*color):
            word += str(run.text)


        elif word != "":
            coloredWords.append(word)
            sentences.append(word)
            parents.append(word)
            word = ""

        # if word == "":
        #     coloredWords.append(word)
        #     sentences.append(word)
        #     parents.append(word)
        #     noChild.append(word)


    report3.save('report3.docx')
#GetOrphanTags()
    if word != "":
        coloredWords.append(word + "\n")
        # word = removeAfter(word)
        child.append(word)
        withChild.append(word)

    return coloredWords


def openFile():
    global filepath
    global filepath2
    filepath = filedialog.askopenfilename(initialdir="/",
                                          title="",
                                          filetypes=(("word documents", "*.docx"),
                                                     ("all files", "*.*")))
    file = open(filepath, 'r')
    # print(filepath)
    file.close()
    filepath2 = str(filepath)
    # filepath2 = '"' + filepath + '"'
    print(filepath2)

    return filepath2


def generateReport():
    fullText = readtxt(filename=filepath2,
                       color=(255, 0, 0))
    # filtered_L = readtxt(filename=filepath2,
    #                   color=(255, 0, 0))
    fullText10 = str(fullText)
    s = ''.join(fullText10)
    w = (s.replace(']', ']\n\n'))
    # w = (w.replace ('\n[', '['))
    # print('\n' + w)
    paragraph = report3.add_paragraph()
    runner = paragraph.add_run("\n" + filepath2 + "\n")
    runner.bold = True  # makes the header bold
    w = (w.replace('([', ''))
    w = (w.replace(',', ''))
    w = (w.replace('' '', ''))

    # print(w)
    # print(fullText)

    table = report3.add_table(rows=1, cols=2)

    # Adding heading in the 1st row of the table
    row = table.rows[0].cells
    row[0].text = 'Parent Tag'
    row[1].text = 'Child Tag/Tags'

    # Adding style to a table
    table.style = 'Colorful List'

    # Now save the document to a location
    report3.save('report3.docx')

    # print(filtered_L)
    # print(fullText)
    # print(fullText2)
    e = 0
    # print(child)
    child2 = removeAfter(child)  # removes everything after the child tag if there is anything to remove
    # print(child2)
    while sentences and child2:
        row = table.add_row().cells  # Adding a row and then adding data in it.
        row[0].text = sentences[0]
        # green = paragraph.add_run(sentences[0] + "    ")
        # paragraph.add_run("\n\n")
        # green.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
        # green.bold = True
        sentences.remove(sentences[0])
        # print(fullText)
        # print(filtered_L)

        if e < len(fullText2):

            if fullText2[e] in filtered_L:

                # row = table.add_row().cells # Adding a row and then adding data in it.
                # print("no child")
                # print("yes")

                row[1].text = " "

                e += 1

            elif fullText2[e] not in filtered_L:
                # print("has a child")
                # print(fullText[e])
                # row[1].text = "Has no child tag"
                if child2:
                    # red = paragraph.add_run(child[0])
                    row[1].text = child2[0]
                    # paragraph.add_run("\n\n")
                    # red.bold = True
                    # red.font.color.rgb = RGBColor(255, 0, 0)
                    child2.remove(child2[0])

                    e += 1

    # while sentences:
    # row = table.add_row().cells # Adding a row and then adding data in it.
    # row[0].text = sentences[0]
    # green = paragraph.add_run(sentences[0])
    # paragraph.add_run("\n\n")
    # green.font.color.rgb = RGBColor(0x00, 0xFF, 0x00)
    # green.bold = True
    # sentences.remove(sentences[0])
    child2.clear()
    sentences.clear()
    child.clear()
    # paragraph.add_run(f)
    report3.save('report3.docx')


def removeAfter(childtags):  # removes everything after the child tag, example "pass"
    seperator = ']'
    childAfter = [i.rsplit(']', 1)[0] + seperator for i in childtags]

    return childAfter


def getDocument():
    if platform.system() == 'Darwin':
        subprocess.check_call(['open', 'report3.docx'])
    elif platform.system() == 'Windows':
        os.startfile('report3.docx')
    # os.startfile(report3) # try either one for windows if the first option gives error
    else:
        subprocess.call('xdg-open', report3)  # for other linux platforms


def getAllParentChild():
    if platform.system() == 'Darwin':
        subprocess.check_call(['open', 'report1.docx'])
    elif platform.system() == 'Windows':
        os.startfile('report1.docx')
    # os.startfile(report3) # try either one for windows if the first option gives error
    else:
        subprocess.call('xdg-open', report3)  # for other linux platforms


if __name__ == '__main__':
    # Creates a word document, saves it as "report 3, and also adds a heading
    report3 = Document()
    report3.add_heading('Report', 0)  # create word document
    paragraph = report3.add_paragraph()
    report3.save('report3.docx')

    # declaring different lists that will be used to store, tags and sentences
    sentences = []
    parent = []
    child = []
    noChild = []
    withChild = []
    parents = []

    # fulltext9 = ''.join(parent)
    # g = str(fulltext9)
    # f = (g.replace (']', ']\n\n'))

    # Creates the gui
    window = Tk(className=' TARGEST')
    # p1 = PhotoImage(file='NorwegianFlag.png')
    # p1 = PIL.Image.open("NorwegianFlag.png")
    # photo = PIL.ImageTk.PhotoImage(p1)
    #
    # label = Label(window, image=photo)
    # label.image = photo
    # label.pack()

    # Setting icon of master window
    # window.iconphoto(False, p1)
    # set window size
    window.geometry("300x300")
    # Creates button 1
    button = Button(text="Choose Document", command=openFile)
    button.pack()
    getAllRelations = Button(window, text="All Tags Relationship", command=getAllParentChild)
    getAllRelations.pack()
    getDoc = Button(window, text="Open Generated Report", command=getDocument)
    getDoc.pack()
    # Creates button 2
    Button(window, text="Generate Report ", command=generateReport).pack()
    # Creates button 3
    button = Button(text="End Program", command=window.destroy)
    button.pack()

    window.mainloop()

    # print(sentences) #parent tags
    # print(child) #child tags
    # print(parent)
    # print(noChild)
    # print(withChild)
    # print(parents)

    # print(fullText)
    # lister2 = [fullText]
    # d = ''.join(fullText)
    # print(fullText)
    # print(d)
    # words = d.split('')
    # words2 = d.split('')
    # print(words)
